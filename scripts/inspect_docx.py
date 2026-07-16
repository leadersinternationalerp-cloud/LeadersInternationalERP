from docx import Document
from pathlib import Path

sample_paths = [
    Path(r'F:/SCHEMES/MATHEMATICS/0096 Mathematics Stage 1 Scheme of Work.docx'),
    Path(r'F:/SCHEMES/ENGLISH/0058 English Stage 1 Scheme of Work.docx'),
    Path(r'F:/SCHEMES/SCIENCE/0097  Primary Science  Stage 1 Scheme of Work.docx')
]

for path in sample_paths:
    print('===', path)
    doc = Document(path)
    print('paragraphs:', len(doc.paragraphs))
    for i, p in enumerate(doc.paragraphs[:120], 1):
        text = p.text.strip()
        if text:
            style = p.style.name if p.style is not None else 'None'
            print(f'{i:03d} [{style}] {text}')
    print('tables:', len(doc.tables))
    for ti, table in enumerate(doc.tables[:3], 1):
        print('TABLE', ti)
        for ri, row in enumerate(table.rows[:6], 1):
            print(' | '.join(cell.text.strip().replace('\n', ' / ') for cell in row.cells))
        print('---')
    print()