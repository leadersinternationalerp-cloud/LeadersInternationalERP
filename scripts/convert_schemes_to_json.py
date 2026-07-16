"""Convert Cambridge scheme-of-work DOCX files to the curriculum JSON used by the app."""

import json
import re
import sys
from collections import OrderedDict
from pathlib import Path

from docx import Document

SCHEMES_ROOT = Path(r"F:/SCHEMES")
OUTPUT_PATH = Path(r"f:/LEADERS_INTERNATIONAL_FILES/LeadersInternationalERP/src/lib/data/curriculum/all_subjects.json")

SUBJECT_MAP = {
    "ART&CRAFT": "Art and Craft", "COMPUTING": "Computing",
    "DIGITAL LITERACY": "Digital Literacy", "ENGLISH": "English Language",
    "GLOBAL PERSPECTIVES": "Global Perspectives", "MATHEMATICS": "Mathematics",
    "MUSIC": "Music", "SCIENCE": "Science",
}
STAGE_RE = re.compile(r"\bStage\s*(\d+)\b", re.I)
UNIT_RE = re.compile(r"^Unit\s+(\d+(?:\.\d+)?[a-z]?)(?:\s+Topic\s+(\d+))?\s*[:\-–—]?\s*(.+?)\s*$", re.I)
OBJECTIVE_CODE_RE = re.compile(r"(?=\b(?:\d{1,2}[A-Za-z]{1,8}|[A-Z]{1,8})\.\d{2}\b)")
BOILERPLATE_RE = re.compile(r"\b(contents|sample lesson|changes to (this )?scheme|copyright|cambridge assessment)\b", re.I)


def clean(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").split()).strip()


def topic_from_text(text: str):
    """Return (unit-id, display title) only for a genuine unit/topic heading."""
    text = clean(text)
    match = UNIT_RE.match(text)
    if not match or BOILERPLATE_RE.search(text):
        return None
    title = match.group(3).strip(" .:-–—")
    # Table-of-contents/schedule headings such as just "Unit 1.1" are not topics.
    if not title or title.lower() in {"and suggested order", "suggested teaching time"}:
        return None
    return match.group(1), match.group(2), title


def objective_lines(text: str):
    """Split a learning-objective cell into its coded objectives, preserving source text."""
    text = clean(text)
    if not text or text.lower() in {"learning objective", "learning objectives"}:
        return []
    parts = [clean(part) for part in OBJECTIVE_CODE_RE.split(text) if clean(part)]
    # Cambridge objective cells use objective codes. Keep a non-coded cell only when it
    # is clearly an objective table entry, rather than silently discarding source text.
    return parts if parts else [text]


def table_texts(table):
    return [[clean(cell.text) for cell in row.cells] for row in table.rows]


def learning_objective_column(rows):
    if not rows:
        return None
    for index, value in enumerate(rows[0]):
        value = value.lower()
        if "learning objective" in value:
            return index
    return None


def parse_docx_topics(doc_path: Path):
    """Return stage and ordered topic records extracted from a scheme DOCX."""
    doc = Document(doc_path)
    stage_match = STAGE_RE.search(doc_path.name)
    stage = f"Stage {stage_match.group(1)}" if stage_match else None
    if stage is None:
        for paragraph in doc.paragraphs:
            stage_match = STAGE_RE.search(clean(paragraph.text))
            if stage_match:
                stage = f"Stage {stage_match.group(1)}"
                break

    topics = OrderedDict()  # normalized title -> {id, topic, lines}
    unit_to_key = {}
    title_to_key = {}
    detailed_unit_ids = set()
    scheduled_units = []

    # Schedule tables are vital for Computing, where objective tables do not repeat a
    # unit heading. Read them first, but retain document order for the final JSON.
    for table in doc.tables:
        for row in table_texts(table)[1:]:
            if row:
                parsed = topic_from_text(row[0])
                if parsed:
                    unit_id, topic_number, title = parsed
                    key = f"{unit_id}:topic:{topic_number}" if topic_number else f"{unit_id}:unit"
                    if key not in topics:
                        topics[key] = {"id": unit_id, "topic": title, "lines": [], "schedule": True}
                        scheduled_units.append(key)
                    title_to_key[title.casefold()] = key
                    if not topic_number:
                        unit_to_key[unit_id] = key

    current_key = None
    schedule_index = 0
    for table in doc.tables:
        rows = table_texts(table)
        if not rows:
            continue
        first_cell = rows[0][0] if rows[0] else ""
        parsed = topic_from_text(first_cell)
        if parsed:
            unit_id, topic_number, title = parsed
            key = (f"{unit_id}:topic:{topic_number}" if topic_number else
                   title_to_key.get(title.casefold(), unit_to_key.get(unit_id, f"{unit_id}:unit")))
            if key not in topics:
                topics[key] = {"id": unit_id, "topic": title, "lines": [], "schedule": False}
                scheduled_units.append(key)
            elif topics[key]["schedule"]:
                # A schedule can combine alternatives; the actual unit heading is the
                # canonical title displayed to users.
                topics[key]["topic"] = title
                topics[key]["schedule"] = False
            title_to_key[title.casefold()] = key
            if topic_number:
                detailed_unit_ids.add(unit_id)
            else:
                unit_to_key[unit_id] = key
            current_key = key
            continue

        objective_column = learning_objective_column(rows)
        continuation = objective_column is None and current_key is not None and rows[0] and OBJECTIVE_CODE_RE.search(rows[0][0])
        if objective_column is None and not continuation:
            continue
        if continuation:
            objective_column = 0
        # When a document has no repeated unit headings (notably Computing), its
        # objective tables follow the unit schedule in sequence.
        if current_key is None:
            while schedule_index < len(scheduled_units) and topics[scheduled_units[schedule_index]]["lines"]:
                schedule_index += 1
            if schedule_index < len(scheduled_units):
                current_key = scheduled_units[schedule_index]
        if current_key is None:
            continue
        record = topics[current_key]
        data_rows = rows if continuation else rows[1:]
        for row in data_rows:
            if objective_column < len(row):
                for line in objective_lines(row[objective_column]):
                    if line not in record["lines"]:
                        record["lines"].append(line)
        has_objectives = bool(record["lines"])
        if has_objectives and current_key in scheduled_units:
            schedule_index = max(schedule_index, scheduled_units.index(current_key) + 1)
        # Do not leak a unit's objective table into a later document section.
        if has_objectives:
            current_key = None

    return stage, [
        {"topic": item["topic"], "content": "\n".join(item["lines"])}
        for item in topics.values()
        if not (item["schedule"] and item["id"] in detailed_unit_ids)
    ]


def build_subject_json():
    data = {subject: {f"Stage {stage}": [] for stage in range(1, 7)} for subject in SUBJECT_MAP.values()}
    failures, missing = [], []
    for folder, subject in SUBJECT_MAP.items():
        directory = SCHEMES_ROOT / folder
        files = sorted(directory.glob("*.docx")) if directory.exists() else []
        found_stages = set()
        for file in files:
            try:
                stage, topics = parse_docx_topics(file)
            except Exception as exc:
                failures.append(f"{file}: {exc}")
                continue
            if stage not in data[subject]:
                failures.append(f"{file}: stage could not be detected")
                continue
            found_stages.add(stage)
            data[subject][stage] = topics
            if not topics or any(not topic["content"].strip() for topic in topics):
                failures.append(f"{file}: parsed {len(topics)} topics; one or more have no learning objectives")
        for number in range(1, 7):
            stage = f"Stage {number}"
            if stage not in found_stages:
                missing.append(f"{subject} {stage}")
    return data, failures, missing


def print_summary(data):
    print(f"{'Subject':<22} {'Stage':<8} {'Topics':>6} {'With content':>13}")
    for subject, stages in data.items():
        for stage, topics in stages.items():
            filled = sum(bool(topic["content"].strip()) for topic in topics)
            print(f"{subject:<22} {stage:<8} {len(topics):>6} {filled:>13}")


if __name__ == "__main__":
    data, failures, missing = build_subject_json()
    print_summary(data)
    if missing:
        print("\nMissing source DOCX files:")
        print("\n".join(f"- {item}" for item in missing))
    if failures:
        print("\nParse failures:", file=sys.stderr)
        print("\n".join(f"- {item}" for item in failures), file=sys.stderr)
        raise SystemExit(1)
    with OUTPUT_PATH.open("w", encoding="utf-8") as output:
        json.dump(data, output, indent=2, ensure_ascii=False)
        output.write("\n")
    print(f"\nWritten {OUTPUT_PATH}")
